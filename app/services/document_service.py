import io
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from core.config import TEMPLATES_DIR


def latex_escape(value: str) -> str:
    escaped = (
        (value or "")
        .replace("\\", r"\textbackslash{}")
        .replace("&", r"\&")
        .replace("%", r"\%")
        .replace("$", r"\$")
        .replace("#", r"\#")
        .replace("_", r"\_")
        .replace("{", r"\{")
        .replace("}", r"\}")
        .replace("~", r"\textasciitilde{}")
        .replace("^", r"\textasciicircum{}")
    )
    return escaped


def split_name(full_name: str) -> tuple[str, str]:
    parts = (full_name or "").strip().split()
    if not parts:
        return "Candidate", ""
    first_name = parts[0]
    last_name = " ".join(parts[1:]) if len(parts) > 1 else ""
    return first_name, last_name


def split_location(location: str) -> tuple[str, str]:
    parts = [segment.strip() for segment in (location or "").split(",") if segment.strip()]
    if not parts:
        return "Not Provided", "Not Provided"
    if len(parts) == 1:
        return parts[0], "Not Provided"
    return ", ".join(parts[:-1]), parts[-1]


def normalize_cover_letter_body(raw_text: str) -> str:
    assert raw_text and raw_text.strip(), "cover_letter body must be non-empty."
    text = raw_text.strip()
    text = re.sub(r"(?is)^dear\s+[^\n,]+,\s*", "", text).strip()
    text = re.sub(
        r"(?is)(kind regards|best regards|sincerely|yours faithfully|yours sincerely)\s*,?\s*[\s\S]*$",
        "",
        text,
    ).strip()
    return text


def build_cover_letter_tex(template_content: str, profile: dict, job_data: dict, cover_letter: dict) -> str:
    first_name, last_name = split_name(profile.get("name", ""))
    address, country = split_location(profile.get("location", ""))

    subject = (cover_letter.get("subject") or job_data.get("title") or "Application").strip()
    company_name = job_data.get("company", "").strip()
    company_location = job_data.get("location", "").strip()
    addressee = "Hiring Manager"
    body = normalize_cover_letter_body(cover_letter.get("body") or "")

    replacements = {
        "{{FIRSTNAME}}": latex_escape(first_name),
        "{{LASTNAME,}}": latex_escape(f"{last_name}," if last_name else ""),
        "{{ADDRESS}}": latex_escape(address),
        "{{COUNTRY}}": latex_escape(country),
        "{{MOBILE}}": latex_escape(profile.get("phone", "")),
        "{{EMAIL}}": latex_escape(profile.get("email", "")),
        "{{COMPANYNAME}}": latex_escape(company_name),
        "{{COMPANYLOCATION}}": latex_escape(company_location),
        "{{SUBJECT}}": latex_escape(subject),
        "{{ADDRESSE}}": latex_escape(addressee),
        "{{COVER_LETTER_CONTENT}}": latex_escape(body),
    }

    rendered = template_content
    for placeholder, value in replacements.items():
        rendered = rendered.replace(placeholder, value)
    return rendered


def build_cover_letter_docx(profile: dict, job_data: dict, cover_letter: dict) -> bytes:
    subject = (cover_letter.get("subject") or job_data.get("title") or "Application").strip()
    body = normalize_cover_letter_body(cover_letter.get("body") or "")
    document = Document()
    if profile.get("name"):
        document.add_heading(profile["name"], level=1)
    document.add_heading(subject, level=2)
    for paragraph in body.split("\n\n"):
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_cv_docx(profile: dict, cv_payload: dict) -> bytes:
    cv_text = (cv_payload.get("cv_text") or "").strip()
    assert cv_text, "cv.cv_text is required to generate CV."
    document = Document()
    if profile.get("name"):
        document.add_heading(profile["name"], level=1)
    document.add_heading("Curriculum Vitae", level=2)
    for line in cv_text.splitlines():
        document.add_paragraph(line)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_cv_pdf(cv_payload: dict) -> bytes:
    cv_text = (cv_payload.get("cv_text") or "").strip()
    assert cv_text, "cv.cv_text is required to generate CV."
    buffer = io.BytesIO()
    pdf_canvas = canvas.Canvas(buffer, pagesize=letter)
    _, page_height = letter
    textobj = pdf_canvas.beginText(40, page_height - 40)
    textobj.setFont("Helvetica", 11)
    for line in cv_text.splitlines():
        if textobj.getY() < 40:
            pdf_canvas.drawText(textobj)
            pdf_canvas.showPage()
            textobj = pdf_canvas.beginText(40, page_height - 40)
            textobj.setFont("Helvetica", 11)
        textobj.textLine(line)
    pdf_canvas.drawText(textobj)
    pdf_canvas.save()
    return buffer.getvalue()


def render_latex_cover_pdf(profile: dict, job_data: dict, cover_letter: dict) -> bytes:
    assert shutil.which("pdflatex"), "pdflatex not found. Please install a LaTeX distribution."
    template_path = TEMPLATES_DIR / "cl_template.tex"
    template_content = template_path.read_text(encoding="utf-8")
    rendered_tex = build_cover_letter_tex(template_content, profile, job_data, cover_letter)

    with tempfile.TemporaryDirectory(prefix="cover_letter_") as temp_dir:
        temp_path = Path(temp_dir)
        tex_path = temp_path / "cover_letter.tex"
        tex_path.write_text(rendered_tex, encoding="utf-8")

        compile_command = [
            "pdflatex",
            "-interaction=nonstopmode",
            "-halt-on-error",
            tex_path.name,
        ]
        run_result = subprocess.run(
            compile_command,
            cwd=temp_path,
            check=False,
            capture_output=True,
            text=True,
        )
        assert run_result.returncode == 0, (
            "Failed to compile cover letter template with pdflatex."
            f"\n{run_result.stdout}\n{run_result.stderr}"
        )
        pdf_path = temp_path / "cover_letter.pdf"
        assert pdf_path.exists(), "PDF compilation succeeded but output file was not created."
        return pdf_path.read_bytes()
