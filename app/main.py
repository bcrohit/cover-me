import base64
import binascii
import io
from pathlib import Path

from docx import Document
from dotenv import load_dotenv
from flask import Flask, jsonify, request
from flask_cors import CORS
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from utils import get_job_data, parse_pdf, save_json
from chat_utils import generate_cover_content, structure_parsed_cv, generate_cv_content
app = Flask(__name__)
CORS(app)
CONTENTS_DIR = Path("assets/contents")

load_dotenv()


def has_profile_content(profile: dict) -> bool:
    return any(bool(value) for value in profile.values())


@app.route("/api/jobdata", methods=["POST"])
def receive_job_data():
    data = request.get_json() or {}
    print("Received profile/job data")
    try:
        assert isinstance(data, dict), "Request payload must be a JSON object."

        profile_mode = data.get("profileMode")
        job_data = data.get("jobData")
        assert job_data is not None, "No job data retrieved."
        save_json("job_data.json", job_data, CONTENTS_DIR)

        if profile_mode == "manual":
            profile = data.get("profile")
            save_json("profile.json", profile, CONTENTS_DIR)
        else:
            encoded_data = data.get("data", "")
            decoded_bytes = base64.b64decode(encoded_data, validate=True)
            assert decoded_bytes.startswith(b"%PDF"), "Uploaded file is not a valid PDF."
            profile = parse_pdf(decoded_bytes)
            profile = structure_parsed_cv(profile)
            save_json("profile.json", profile, CONTENTS_DIR)

        cover_letter = generate_cover_content(job_data, profile)
        cv_content = generate_cv_content(job_data, profile)
        save_json("cover_letter.json", cover_letter, CONTENTS_DIR)
        save_json("cv.json", cv_content, CONTENTS_DIR)
        return jsonify(
            {
                "status": "success",
                "message": "Cover letter generated.",
                "cover_letter": cover_letter,
                "cv_content": cv_content,
            }
        )
    except AssertionError as e:
        print("Assertion error:", e)
        return jsonify({"status": "error", "message": str(e)}), 400
    except binascii.Error:
        return jsonify({"status": "error", "message": "Invalid base64 payload."}), 400
    except Exception as e:
        print("Error saving data:", e)
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/generate", methods=["POST"])
def generate_documents():
    """Generates a preview, DOCX and PDF from jobData and profile and returns files as base64."""
    payload = request.get_json() or {}
    # Accept either { jobData, profile } or read saved files
    profile = payload.get("profile") or {}
    if payload.get("profileMode") == "upload":
        profile = {}

    try:
        if not profile:
            profile_path = CONTENTS_DIR / "profile.json"
            if profile_path.exists():
                profile = get_job_data("profile.json", CONTENTS_DIR)

        # Build a simple preview text (cover letter + CV summary)
        name = profile.get("name", "")
        skills = profile.get("skills", "")
        experience = profile.get("experience", "")
        projects = profile.get("projects", "")

        cover_letter_text = get_job_data("cover_letter.json", CONTENTS_DIR).get("cover_letter", "")
        assert cover_letter_text, "cover_letter.json must include a non-empty cover_letter."
        cover_lines = cover_letter_text.split("\n")
        preview = "\n".join(cover_lines)

        # Create DOCX
        doc = Document()
        if name:
            doc.add_heading(name, level=1)
        doc.add_heading("Cover Letter", level=2)
        for line in cover_lines:
            doc.add_paragraph(line)

        doc.add_page_break()
        doc.add_heading("Curriculum Vitae", level=2)
        if experience:
            doc.add_heading("Experience", level=3)
            doc.add_paragraph(experience)
        if projects:
            doc.add_heading("Projects", level=3)
            for p in [s.strip() for s in projects.split(",") if s.strip()]:
                doc.add_paragraph(p, style="List Bullet")
        if skills:
            doc.add_heading("Skills", level=3)
            doc.add_paragraph(skills)

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()
        docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")

        # Create a simple PDF using reportlab
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        width, height = letter
        textobj = c.beginText(40, height - 40)
        textobj.setFont("Helvetica", 11)
        for paragraph in cover_lines:
            # Simple wrapping
            for line in paragraph.split("\n"):
                textobj.textLine(line)
            textobj.textLine("")
        c.drawText(textobj)
        c.showPage()

        # Append basic CV lines
        textobj = c.beginText(40, height - 40)
        textobj.setFont("Helvetica-Bold", 12)
        textobj.textLine("Curriculum Vitae")
        textobj.textLine("")
        textobj.setFont("Helvetica", 11)
        if experience:
            textobj.textLine("Experience:")
            for experience_line in experience.split("\n"):
                textobj.textLine("- " + experience_line)
            textobj.textLine("")
        if projects:
            textobj.textLine("Projects:")
            for p in [s.strip() for s in projects.split(",") if s.strip()]:
                textobj.textLine("- " + p)
            textobj.textLine("")
        if skills:
            textobj.textLine("Skills: " + skills)

        c.drawText(textobj)
        c.showPage()
        c.save()

        pdf_bytes = pdf_buffer.getvalue()
        pdf_b64 = base64.b64encode(pdf_bytes).decode("utf-8")

        files = [
            {
                "filename": f"{(name or 'candidate').replace(' ', '_')}_cover_and_cv.docx",
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "data": docx_b64,
            },
            {
                "filename": f"{(name or 'candidate').replace(' ', '_')}_cover_and_cv.pdf",
                "content_type": "application/pdf",
                "data": pdf_b64,
            },
        ]

        return jsonify({"preview": preview, "files": files})
    except Exception as e:
        print("Generation error:", e)
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
