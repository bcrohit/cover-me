from flask import Flask, request, jsonify
from flask_cors import CORS
import json
import io
import base64
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from utils import get_job_data

app = Flask(__name__)
CORS(app)

@app.route("/api/jobdata", methods=["POST"])
def receive_job_data():
    data = request.get_json()
    print("Received Job Data")
    # Accept either raw jobData or { jobData, profile }
    try:
        if isinstance(data, dict) and 'jobData' in data:
            job = data.get('jobData')
            profile = data.get('profile', {})
        else:
            # legacy: assume whole payload is job data
            job = data
            profile = {}

        # save job and profile separately for clarity
        with open("assets/contents/job_data.json", "w", encoding='utf-8') as f:
            json.dump(job, f, ensure_ascii=False, indent=2)
        with open("assets/contents/profile.json", "w", encoding='utf-8') as f:
            json.dump(profile, f, ensure_ascii=False, indent=2)
        return jsonify({"status": "success", "message": "Data received"})
    except Exception as e:
        print("Error saving data:", e)
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/generate", methods=["POST"])
def generate_documents():
    """Generates a preview, DOCX and PDF from jobData and profile and returns files as base64."""
    payload = request.get_json() or {}
    # Accept either { jobData, profile } or read saved files
    job = payload.get('jobData') or {}
    profile = payload.get('profile') or {}

    try:
        # Build a simple preview text (cover letter + CV summary)
        name = profile.get('name', '')
        skills = profile.get('skills', '')
        experience = profile.get('experience', '')
        projects = profile.get('projects', '')
        title = job.get('title', '')
        company = job.get('company', '')

        cover_lines = get_job_data("output.json").get("cover_letter", "").split("\n")
        preview = "\n".join(cover_lines)

        # Create DOCX
        doc = Document()
        if name:
            doc.add_heading(name, level=1)
        doc.add_heading('Cover Letter', level=2)
        for line in cover_lines:
            doc.add_paragraph(line)

        doc.add_page_break()
        doc.add_heading('Curriculum Vitae', level=2)
        if experience:
            doc.add_heading('Experience', level=3)
            doc.add_paragraph(experience)
        if projects:
            doc.add_heading('Projects', level=3)
            for p in [s.strip() for s in projects.split(',') if s.strip()]:
                doc.add_paragraph(p, style='List Bullet')
        if skills:
            doc.add_heading('Skills', level=3)
            doc.add_paragraph(skills)

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()
        docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')

        # Create a simple PDF using reportlab
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        width, height = letter
        textobj = c.beginText(40, height - 40)
        textobj.setFont('Helvetica', 11)
        for paragraph in cover_lines:
            # Simple wrapping
            for line in paragraph.split('\n'):
                textobj.textLine(line)
            textobj.textLine('')
        c.drawText(textobj)
        c.showPage()

        # Append basic CV lines
        textobj = c.beginText(40, height - 40)
        textobj.setFont('Helvetica-Bold', 12)
        textobj.textLine('Curriculum Vitae')
        textobj.textLine('')
        textobj.setFont('Helvetica', 11)
        if experience:
            textobj.textLine('Experience:')
            for l in experience.split('\n'):
                textobj.textLine('- ' + l)
            textobj.textLine('')
        if projects:
            textobj.textLine('Projects:')
            for p in [s.strip() for s in projects.split(',') if s.strip()]:
                textobj.textLine('- ' + p)
            textobj.textLine('')
        if skills:
            textobj.textLine('Skills: ' + skills)

        c.drawText(textobj)
        c.showPage()
        c.save()

        pdf_bytes = pdf_buffer.getvalue()
        pdf_b64 = base64.b64encode(pdf_bytes).decode('utf-8')

        files = [
            {
                'filename': f'{(name or "candidate").replace(" ","_")}_cover_and_cv.docx',
                'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'data': docx_b64
            },
            {
                'filename': f'{(name or "candidate").replace(" ","_")}_cover_and_cv.pdf',
                'content_type': 'application/pdf',
                'data': pdf_b64
            }
        ]

        return jsonify({'preview': preview, 'files': files})
    except Exception as e:
        print('Generation error:', e)
        return jsonify({'error': str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
